#!/usr/bin/env python3
"""
Generate a DOCX from a cv-variant.md (or cv.md) file.

Usage:
    python3 scripts/generate-cv-docx.py opportunities/safetica-senior-pm/cv-variant.md

Output is saved alongside the source file as cv-<firstname>-<company>.docx
If the source is context/cv.md, output goes to context/cv-<firstname>.docx
"""

import sys
import os

# Reuse the parser from generate-cv-pdf.py
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from generate_cv_pdf_parser import parse_cv_markdown, determine_output_path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Styles and helpers
# ---------------------------------------------------------------------------

FONT_NAME = "Calibri"
COLOR_ACCENT = RGBColor(0x37, 0x41, 0x51)
COLOR_SUBTLE = RGBColor(0x6B, 0x72, 0x80)
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_BLACK = RGBColor(0x11, 0x11, 0x11)


def set_run_font(run, size_pt, color=COLOR_BODY, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, 9, COLOR_ACCENT, bold=True)
    run.font.letter_spacing = Pt(0.8)
    # Add bottom border via paragraph format
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    return p


def _remove_table_borders(table):
    """Remove all borders from a table."""
    from docx.oxml.ns import qn
    from lxml import etree
    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl, qn("w:tblPr"))
    borders = etree.SubElement(tbl_pr, qn("w:tblBorders"))
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = etree.SubElement(borders, qn(f"w:{border_name}"))
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")


def add_job_header(doc, title, dates, company_line):
    # Title + dates on same line using a table for alignment
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.columns[0].width = Inches(4.2)
    table.columns[1].width = Inches(2.3)
    _remove_table_borders(table)

    # Title cell
    p_title = table.rows[0].cells[0].paragraphs[0]
    p_title.space_before = Pt(0)
    p_title.space_after = Pt(0)
    run = p_title.add_run(title)
    set_run_font(run, 9.5, COLOR_BLACK, bold=True)

    # Dates cell
    p_dates = table.rows[0].cells[1].paragraphs[0]
    p_dates.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dates.space_before = Pt(0)
    p_dates.space_after = Pt(0)
    run = p_dates.add_run(dates)
    set_run_font(run, 8, COLOR_SUBTLE, bold=True)

    # Company line
    p_company = doc.add_paragraph()
    p_company.space_before = Pt(0)
    p_company.space_after = Pt(2)
    run = p_company.add_run(company_line)
    set_run_font(run, 8.5, COLOR_SUBTLE)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.space_before = Pt(0.5)
    p.space_after = Pt(1)
    # Clear default run and add our own
    p.clear()
    run = p.add_run(text)
    set_run_font(run, 8.5, COLOR_BODY)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.2)
    pf.first_line_indent = Inches(-0.15)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_docx(data):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(9)
    style.font.color.rgb = COLOR_BODY

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    # --- Header: Name ---
    p_name = doc.add_paragraph()
    p_name.space_after = Pt(2)
    run = p_name.add_run(data["name"])
    set_run_font(run, 22, COLOR_BLACK, bold=True)

    # --- Contact line ---
    contact_parts = []
    if data["contact"].get("location"):
        contact_parts.append(data["contact"]["location"])
    if data["contact"].get("email"):
        contact_parts.append(data["contact"]["email"])
    if data["contact"].get("linkedin"):
        contact_parts.append(data["contact"]["linkedin"])
    if data["contact"].get("phone"):
        contact_parts.append(data["contact"]["phone"])

    p_contact = doc.add_paragraph()
    p_contact.space_after = Pt(6)
    run = p_contact.add_run("  ·  ".join(contact_parts))
    set_run_font(run, 8.5, COLOR_SUBTLE)

    # --- Summary ---
    if data["summary"]:
        p_summary = doc.add_paragraph()
        p_summary.space_before = Pt(4)
        p_summary.space_after = Pt(8)
        run = p_summary.add_run(data["summary"])
        set_run_font(run, 9, RGBColor(0x37, 0x41, 0x51))
        pf = p_summary.paragraph_format
        pf.left_indent = Inches(0.1)
        pf.right_indent = Inches(0.1)

    # --- Personal Project ---
    if data.get("personal_project"):
        add_section_heading(doc, "Personal Project (In Progress)")
        for proj in data["personal_project"]:
            p_proj_title = doc.add_paragraph()
            p_proj_title.space_before = Pt(2)
            p_proj_title.space_after = Pt(0)
            run = p_proj_title.add_run(proj["title"])
            set_run_font(run, 9.5, COLOR_BLACK, bold=True)

            if proj.get("subtitle"):
                p_sub = doc.add_paragraph()
                p_sub.space_before = Pt(0)
                p_sub.space_after = Pt(2)
                run = p_sub.add_run(proj["subtitle"])
                set_run_font(run, 8.5, COLOR_SUBTLE)

            for achievement in proj["achievements"]:
                add_bullet(doc, achievement)

    # --- Experience ---
    add_section_heading(doc, "Experience")
    for role in data["experience"]:
        company_parts = [role["company"]]
        if role.get("location"):
            company_parts.append(role["location"])
        if role.get("note"):
            company_parts.append(role["note"])
        company_line = " · ".join(company_parts)

        add_job_header(doc, role["title"], role["dates"], company_line)
        for achievement in role["achievements"]:
            add_bullet(doc, achievement)

    # --- Skills ---
    add_section_heading(doc, "Skills")
    for skill in data["skills"]:
        p = doc.add_paragraph()
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        run_label = p.add_run(f"{skill['category']}: ")
        set_run_font(run_label, 8, COLOR_ACCENT, bold=True)
        run_value = p.add_run(skill["items"])
        set_run_font(run_value, 8.5, RGBColor(0x44, 0x44, 0x44))

    # --- Education ---
    add_section_heading(doc, "Education")
    for edu in data["education"]:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.columns[0].width = Inches(4.2)
        table.columns[1].width = Inches(2.3)
        _remove_table_borders(table)

        p_left = table.rows[0].cells[0].paragraphs[0]
        p_left.space_before = Pt(0)
        p_left.space_after = Pt(0)
        run = p_left.add_run(edu["degree"])
        set_run_font(run, 8.5, COLOR_BLACK, bold=True)
        run = p_left.add_run(f" — {edu['school']}")
        set_run_font(run, 8.5, COLOR_SUBTLE)

        p_right = table.rows[0].cells[1].paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.space_before = Pt(0)
        p_right.space_after = Pt(0)
        run = p_right.add_run(edu["dates"])
        set_run_font(run, 8, COLOR_SUBTLE, bold=True)

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/generate-cv-docx.py <path-to-cv-variant.md>")
        sys.exit(1)

    source_path = sys.argv[1]
    if not os.path.isfile(source_path):
        print(f"ERROR: File not found: {source_path}")
        sys.exit(1)

    with open(source_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    data = parse_cv_markdown(md_content)
    output_path = determine_output_path(source_path, data["name"]).replace(".pdf", ".docx")

    doc = build_docx(data)
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")


if __name__ == "__main__":
    main()
